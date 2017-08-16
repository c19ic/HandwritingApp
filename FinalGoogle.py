#Code will allow user to upload a text file into Google Drive
#This is the final Google Code, for additional possibilities, refer back to Google1

#Imports
import webbrowser
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from docx import Document

#Authentication process where Google asks user to sign in
gauth = GoogleAuth()        
gauth.LocalWebserverAuth()
drive = GoogleDrive(gauth)


#Open a text file and gets its strings into a word document
file_object = open("C:/Python27/Stem/Documents/Document1.txt","r")
x= file_object.read()

document = Document()
document.add_heading("Edit your Title",0)
document.add_paragraph(x)
document.save("Doc1.docx")

file_object.close()

#Beginning of uploading to Google Drive
webbrowser.open('https://drive.google.com/drive/my-drive') #Opens home page for Google Drive
file1 = drive.CreateFile() #Creates a new file
file1.SetContentFile('Doc1.docx') #Gets the text file called Document1
file1.Upload()  #Uploads file into Google Drive


